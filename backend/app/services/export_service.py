from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List

from app.models import Analysis, Opportunity

class ExportService:
    """Service for exporting analysis results to DOCX format."""
    
    def generate_analysis_docx(self, analysis: Analysis, opportunity: Opportunity) -> BytesIO:
        """
        Generate a professional DOCX document from analysis results.
        Returns a BytesIO buffer containing the DOCX file.
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Gap Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f'{opportunity.title}\n{opportunity.solicitation_number}')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        
        doc.add_page_break()
        
        # 1. Executive Summary
        doc.add_heading('1. Executive Summary', 1)
        
        # Overall Relevance
        doc.add_heading('Overall Relevance Assessment', 2)
        relevance_para = doc.add_paragraph()
        
        # Display score and label
        score_val = analysis.overall_relevance_score or "0"
        label_val = analysis.overall_relevance_label or "UNKNOWN"
        
        relevance_run = relevance_para.add_run(f'{score_val}% - {label_val}')
        relevance_run.bold = True
        relevance_run.font.size = Pt(14)
        self._apply_relevance_color(relevance_run, label_val)
        
        # Evaluator Perspective
        if analysis.evaluator_perspective:
            doc.add_heading('Evaluator Perspective', 3)
            doc.add_paragraph(analysis.evaluator_perspective)
            
        # Go/No-Go
        doc.add_heading('Recommendation', 2)
        go_no_go_para = doc.add_paragraph()
        go_no_go_run = go_no_go_para.add_run(analysis.go_no_go_recommendation or "N/A")
        go_no_go_run.bold = True
        go_no_go_run.font.size = Pt(14)
        self._apply_go_no_go_color(go_no_go_run, analysis.go_no_go_recommendation)
        
        if analysis.go_no_go_reasoning:
            doc.add_paragraph(analysis.go_no_go_reasoning)
            
        # Red Flags
        if analysis.red_flags:
            doc.add_heading('Red Flags', 2)
            for flag in analysis.red_flags:
                para = doc.add_paragraph(style='List Bullet')
                flag_run = para.add_run(f"🚩 {flag.get('warning', 'Warning')}")
                flag_run.bold = True
                flag_run.font.color.rgb = RGBColor(239, 68, 68) # Red
                doc.add_paragraph(flag.get('reason', ''), style='List Bullet 2')
        
        doc.add_page_break()
        
        # 2. Requirements Compliance Matrix
        if analysis.requirements_matrix:
            doc.add_heading('2. Requirements Compliance Matrix', 1)
            
            summary = analysis.requirements_summary or {}
            if summary:
                summary_para = doc.add_paragraph()
                summary_para.add_run(f"Coverage: {summary.get('coverage_percentage', 0)}% ").bold = True
                summary_para.add_run(f"({summary.get('strong', 0)} Strong, {summary.get('moderate', 0)} Moderate, {summary.get('weak', 0)} Weak, {summary.get('gap', 0)} Gaps)")

            table = doc.add_table(rows=len(analysis.requirements_matrix) + 1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID / Requirement'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Evidence/Notes'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for idx, req in enumerate(analysis.requirements_matrix):
                row = table.rows[idx + 1]
                row.cells[0].text = f"{req.get('req_id', '')}: {req.get('requirement_text', '')}"
                
                status = req.get('coverage_status', 'GAP')
                status_run = row.cells[1].paragraphs[0].add_run(status)
                self._apply_coverage_color(status_run, status)
                
                evidence = req.get('supporting_evidence', [])
                notes = req.get('notes', '')
                evidence_text = "\n".join(evidence) if isinstance(evidence, list) else str(evidence)
                if notes:
                    evidence_text += f"\nNote: {notes}"
                row.cells[2].text = evidence_text
            
            doc.add_page_break()
        
        # 3. Opportunity Overview
        doc.add_heading('3. Opportunity Overview', 1)
        
        overview_data = [
            ('Title', opportunity.title),
            ('Solicitation Number', opportunity.solicitation_number),
            ('Agency', opportunity.agency),
            ('Response Due Date', str(opportunity.response_due_date) if opportunity.response_due_date else 'N/A'),
            ('Estimated Value', str(opportunity.estimated_value) if opportunity.estimated_value else 'N/A'),
            ('NAICS Code', opportunity.naics_code or 'N/A'),
        ]
        
        table = doc.add_table(rows=len(overview_data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for idx, (label, value) in enumerate(overview_data):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value) if value else 'N/A'
        
        doc.add_page_break()
        
        # 4. Dimensional Scores
        doc.add_heading('4. Dimensional Scores', 1)
        
        scores = analysis.dimensional_scores or {}
        
        dimensions = [
            ('Scope Alignment', scores.get('scope_alignment', {})),
            ('Magnitude', scores.get('magnitude', {})),
            ('Complexity', scores.get('complexity', {})),
            ('Recency', scores.get('recency', {})),
            ('Quality', scores.get('quality', {})),
        ]
        
        for name, data in dimensions:
            score = data.get('score', 0)
            label = data.get('label', 'UNKNOWN')
            
            para = doc.add_heading(f'{name}: {score}% ({label})', 3)
            # Find the run in the heading to apply color
            if para.runs:
                self._apply_dimensional_color(para.runs[0], label)
            
            if data.get('strengths'):
                s_para = doc.add_paragraph()
                s_para.add_run('Strengths:').bold = True
                for s in data['strengths']:
                    doc.add_paragraph(f"• {s.get('item', '')}: {s.get('evidence', '')}", style='List Bullet 2')
            
            if data.get('weaknesses'):
                w_para = doc.add_paragraph()
                w_para.add_run('Weaknesses:').bold = True
                for w in data['weaknesses']:
                    doc.add_paragraph(f"• {w.get('item', '')}: {w.get('evidence', '')}", style='List Bullet 2')

            if data.get('gaps'):
                g_para = doc.add_paragraph()
                g_para.add_run('Gaps:').bold = True
                for g in data['gaps']:
                    doc.add_paragraph(f"• {g.get('item', '')}: {g.get('evidence', '')}", style='List Bullet 2')
        
        doc.add_page_break()
        
        # 4. Gap Matrix
        if analysis.gap_matrix and isinstance(analysis.gap_matrix, dict):
            doc.add_heading('4. Gap Matrix', 1)
            
            for category, requirements in analysis.gap_matrix.items():
                doc.add_heading(category, 2)
                
                if requirements and len(requirements) > 0:
                    # Create table: Requirement | Support | Coverage
                    table = doc.add_table(rows=len(requirements) + 1, cols=3)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Requirement'
                    hdr_cells[1].text = 'Supporting References'
                    hdr_cells[2].text = 'Coverage'
                    for cell in hdr_cells:
                        cell.paragraphs[0].runs[0].bold = True
                    
                    # Data
                    for idx, req in enumerate(requirements):
                        row = table.rows[idx + 1]
                        row.cells[0].text = f"{req.get('requirement_id', 'N/A')}: {req.get('text', 'N/A')}"
                        row.cells[1].text = ', '.join(req.get('support_references', [])) or 'None'
                        coverage = req.get('coverage', 'N/A')
                        row.cells[2].text = coverage
                        # Apply color to coverage cell
                        self._apply_coverage_color(row.cells[2].paragraphs[0].runs[0], coverage)
        
        doc.add_page_break()
        
        # 5. Contract-Level Assessments
        if analysis.contract_assessments:
            doc.add_heading('5. Past Performance Portfolio Assessment', 1)
            
            for ca in analysis.contract_assessments:
                doc.add_heading(f"Contract: {ca.get('contract_name', 'N/A')}", 2)
                
                details = [
                    ('Contract #', ca.get('contract_number')),
                    ('Customer', f"{ca.get('customer_agency', 'N/A')} ({ca.get('service_branch', 'N/A')})"),
                    ('Value', str(ca.get('contract_value', 'N/A'))),
                    ('Relevance Score', f"{ca.get('relevance_score', 0)}%"),
                    ('Scope Match', ca.get('scope_match', 'N/A')),
                ]
                
                table = doc.add_table(rows=len(details), cols=2)
                table.style = 'Light Grid Accent 1'
                for i, (label, val) in enumerate(details):
                    row = table.rows[i]
                    row.cells[0].text = label
                    row.cells[0].paragraphs[0].runs[0].bold = True
                    row.cells[1].text = str(val) if val else 'N/A'
                
                if ca.get('limitations'):
                    para = doc.add_paragraph()
                    para.add_run('Limitations: ').bold = True
                    para.add_run(", ".join(ca['limitations']))
                
                doc.add_paragraph(f"Primary Use: {ca.get('primary_use', 'N/A')}")
            
            doc.add_page_break()

        # 6. Strengths
        if analysis.strengths and len(analysis.strengths) > 0:
            doc.add_heading('6. Strengths', 1)
            
            for strength in analysis.strengths:
                doc.add_heading(strength.get('title', 'Strength'), 3)
                doc.add_paragraph(strength.get('description', ''))
                if strength.get('impact'):
                    impact_para = doc.add_paragraph()
                    impact_para.add_run('Impact: ').italic = True
                    impact_para.add_run(str(strength['impact']))
        
        doc.add_page_break()
        
        # 7. Weaknesses & Gaps
        if analysis.weaknesses and len(analysis.weaknesses) > 0:
            doc.add_heading('7. Weaknesses & Gaps', 1)
            
            for weakness in analysis.weaknesses:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(weakness.get('description', '')).bold = True
                
                risk_level = weakness.get('risk_level', 'UNKNOWN')
                risk_para = doc.add_paragraph(style='List Bullet 2')
                risk_run = risk_para.add_run(f'Risk Level: {risk_level}')
                self._apply_risk_color(risk_run, risk_level)
                
                if weakness.get('mitigation'):
                    mitigation_para = doc.add_paragraph(style='List Bullet 2')
                    mitigation_para.add_run('Mitigation: ').italic = True
                    mitigation_para.add_run(weakness['mitigation'])
        
        doc.add_page_break()
        
        # 8. Recommendations
        doc.add_heading('8. Strategic Recommendations', 1)
        
        if analysis.recommendations and isinstance(analysis.recommendations, dict):
            if analysis.recommendations.get('narrative_strategy'):
                doc.add_heading('Narrative Strategy', 2)
                doc.add_paragraph(analysis.recommendations['narrative_strategy'])
            
            # v2.0 gap_mitigations is a list
            mitigations = analysis.recommendations.get('gap_mitigations', [])
            if mitigations:
                doc.add_heading('Gap Mitigation Actions', 2)
                for m in mitigations:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"[{m.get('priority', 'MEDIUM')}] {m.get('gap', '')}: ").bold = True
                    para.add_run(m.get('action', ''))
            elif analysis.recommendations.get('gap_mitigation'):
                # Backward compatibility
                doc.add_heading('Gap Mitigation', 2)
                doc.add_paragraph(analysis.recommendations['gap_mitigation'])
        
        doc.add_page_break()
        
        # 9. Appendix: Documents Analyzed
        doc.add_heading('9. Appendix: Documents Analyzed', 1)
        
        if analysis.documents_analyzed:
            for idx, doc_id in enumerate(analysis.documents_analyzed, 1):
                doc.add_paragraph(f'{idx}. Document ID: {doc_id}', style='List Number')
        else:
            doc.add_paragraph('No documents listed.')
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _apply_relevance_color(self, run, label: str):
        """Apply color based on relevance label."""
        if not label:
            return
        label = label.upper()
        if "VERY RELEVANT" in label:
            run.font.color.rgb = RGBColor(34, 197, 94)  # Green
        elif "RELEVANT" in label:
            run.font.color.rgb = RGBColor(59, 130, 246)  # Blue
        elif "SOMEWHAT" in label:
            run.font.color.rgb = RGBColor(234, 179, 8)  # Yellow
        else:
            run.font.color.rgb = RGBColor(239, 68, 68)  # Red
    
    def _apply_go_no_go_color(self, run, recommendation: str):
        """Apply color based on go/no-go."""
        if not recommendation:
            return
        rec = recommendation.upper()
        if "GO" in rec and "NO" not in rec:
            run.font.color.rgb = RGBColor(34, 197, 94)  # Green
        elif "CONDITIONAL" in rec:
            run.font.color.rgb = RGBColor(234, 179, 8)  # Yellow
        else:
            run.font.color.rgb = RGBColor(239, 68, 68)  # Red
    
    def _apply_dimensional_color(self, run, label: str):
        """Apply color based on dimensional label."""
        if not label:
            return
        label = label.upper()
        if label == "HIGH":
            run.font.color.rgb = RGBColor(34, 197, 94)
        elif label == "MODERATE" or label == "MEDIUM":
            run.font.color.rgb = RGBColor(59, 130, 246)
        else:
            run.font.color.rgb = RGBColor(239, 68, 68)
    
    def _apply_coverage_color(self, run, coverage: str):
        """Apply color based on gap matrix coverage."""
        if not coverage:
            return
        coverage = coverage.upper()
        if coverage == "STRONG":
            run.font.color.rgb = RGBColor(34, 197, 94)
        elif coverage == "MODERATE":
            run.font.color.rgb = RGBColor(59, 130, 246)
        elif coverage == "WEAK":
            run.font.color.rgb = RGBColor(234, 179, 8)
        else:
            run.font.color.rgb = RGBColor(239, 68, 68)
    
    def _apply_risk_color(self, run, risk_level: str):
        """Apply color based on risk level."""
        if not risk_level:
            return
        risk_level = risk_level.upper()
        if risk_level == "HIGH" or risk_level == "CRITICAL":
            run.font.color.rgb = RGBColor(239, 68, 68)
        elif risk_level == "MEDIUM" or risk_level == "MODERATE":
            run.font.color.rgb = RGBColor(234, 179, 8)
        else:
            run.font.color.rgb = RGBColor(34, 197, 94)
