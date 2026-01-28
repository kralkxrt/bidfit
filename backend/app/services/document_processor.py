import io
import re
from typing import Optional, Dict, Any, List
from uuid import UUID
from datetime import datetime
import PyPDF2
from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.services.llm_service import LLMService
from app.services.storage_service import StorageService
from app.services.pdf_text_service import PdfTextService
from app.models import Document, Company, OpportunityDocument


class DocumentProcessor:
    """
    Handles document upload, text extraction, metadata parsing.
    """
    
    def __init__(
        self,
        llm_service: LLMService,
        storage_service: StorageService
    ):
        self.llm = llm_service
        self.storage = storage_service
        self.pdf_text = PdfTextService()
    
    async def process_document(
        self,
        db: AsyncSession,
        file_content: bytes,
        filename: str,
        company_id: UUID,
        document_type: str,
        user_id: UUID
    ) -> Document:
        """
        Process an uploaded document through the full pipeline.
        
        Steps:
        1. Upload file to storage
        2. Extract raw text
        3. Extract metadata using LLM
        4. Create database record
        """
        
        # 1. Upload to storage
        # Need str company_id for storage path
        file_path = await self.storage.upload_file(
            file_content=file_content,
            filename=filename,
            company_id=str(company_id)
        )
        
        # 2. Extract raw text
        mime_type = self._detect_mime_type(filename)
        raw_text = self._extract_text(file_content, mime_type)
        
        # 3. Extract metadata using LLM
        parsed_content = await self._extract_metadata(raw_text, document_type, filename)
        performance_ratings = parsed_content.get("performance_ratings") or {}
        
        # 4. Create document record
        document = Document(
            company_id=company_id,
            document_type=document_type,
            filename=filename,
            file_path=file_path,
            file_size_bytes=len(file_content),
            mime_type=mime_type,
            raw_text=raw_text,
            parsed_content=parsed_content,
            processing_status="completed", 
            processed_at=datetime.utcnow(),
            created_by=user_id,
            
            # Map extracted fields
            contract_number=parsed_content.get("contract_number"),
            contract_title=parsed_content.get("contract_title"),
            customer_agency=parsed_content.get("customer_agency"),
            customer_command=parsed_content.get("customer_command"),
            contract_value=self._parse_currency(parsed_content.get("contract_value")),
            period_of_performance_start=self._parse_date(parsed_content.get("pop_start")),
            period_of_performance_end=self._parse_date(parsed_content.get("pop_end")),
            naics_code=parsed_content.get("naics_code"),
            clearance_level=parsed_content.get("clearance_level"),
            fte_count=self._parse_int(parsed_content.get("fte_count")),
            geographic_scope=parsed_content.get("geographic_scope"),
            locations=self._parse_locations(parsed_content.get("locations")),
            
            # Map CPARS ratings into dedicated columns (if present)
            cpars_quality_rating=self._get_cpars_rating(
                performance_ratings,
                ["quality", "quality of product", "quality of product/service"]
            ),
            cpars_schedule_rating=self._get_cpars_rating(
                performance_ratings,
                ["schedule"]
            ),
            cpars_cost_rating=self._get_cpars_rating(
                performance_ratings,
                ["cost", "cost control"]
            ),
            cpars_management_rating=self._get_cpars_rating(
                performance_ratings,
                ["management", "business relations", "management/business relations"]
            ),
            cpars_overall_rating=self._get_cpars_rating(
                performance_ratings,
                ["overall", "overall rating", "overall performance"]
            )
        )
        
        db.add(document)
        await db.commit()
        await db.refresh(document)

        # 5. Extract and store PDF text positions for citation highlighting
        if mime_type == "application/pdf":
            try:
                positions = self.pdf_text.extract_text_positions(file_content)
                await self.pdf_text.store_text_positions(db, str(document.id), positions)
            except Exception as e:
                print(f"PDF position extraction failed: {e}")
        
        return document

    async def process_opportunity_document(
        self,
        db: AsyncSession,
        file_content: bytes,
        filename: str,
        opportunity_id: UUID,
        document_type: str,
        company_id: UUID
    ) -> OpportunityDocument:
        """
        Process an uploaded opportunity document (SOW/PWS).
        """
        
        # 1. Upload to storage
        # Construct path for opportunity: company_id/opportunities/opp_id/filename
        # We leverage the storage service logic by passing the folder path as company_id
        folder_path = f"{str(company_id)}/opportunities/{str(opportunity_id)}"
        
        file_path = await self.storage.upload_file(
            file_content=file_content,
            filename=filename,
            company_id=folder_path
        )
        
        # 2. Extract raw text
        mime_type = self._detect_mime_type(filename)
        raw_text = self._extract_text(file_content, mime_type)
        
        # 3. Extract Requirements using LLM
        parsed_requirements = await self.llm.extract_requirements(raw_text)
        
        # 4. Create record
        doc = OpportunityDocument(
            opportunity_id=opportunity_id,
            document_type=document_type,
            filename=filename,
            file_path=file_path,
            file_size_bytes=len(file_content),
            mime_type=mime_type,
            raw_text=raw_text,
            parsed_requirements=parsed_requirements,
            processing_status="completed",
            processed_at=datetime.utcnow()
        )
        
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        # 5. Extract and store PDF text positions for citation highlighting
        if mime_type == "application/pdf":
            try:
                positions = self.pdf_text.extract_text_positions(file_content)
                await self.pdf_text.store_opportunity_text_positions(db, str(doc.id), positions)
            except Exception as e:
                print(f"PDF position extraction failed: {e}")
        
        return doc

    async def reprocess_requirements(
        self,
        db: AsyncSession,
        document: OpportunityDocument
    ) -> OpportunityDocument:
        """
        Re-run requirement extraction on an existing document using stored raw_text.
        """
        if not document.raw_text:
            raise ValueError("Document has no raw text to process")
            
        # Extract Requirements using LLM
        parsed_requirements = await self.llm.extract_requirements(document.raw_text)
        
        # Update record
        document.parsed_requirements = parsed_requirements
        document.processed_at = datetime.utcnow()
        document.processing_status = "completed"
        
        db.add(document)
        await db.commit()
        await db.refresh(document)
        
        return document

    
    def _extract_text(self, file_content: bytes, mime_type: str) -> str:
        """Extract plain text from document based on mime type."""
        
        text = ""
        if mime_type == "application/pdf":
            text = self._extract_pdf_text(file_content)
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            text = self._extract_docx_text(file_content)
        elif mime_type == "text/plain":
            text = file_content.decode("utf-8", errors='ignore')
        
        return self._clean_text(text)

    def _clean_text(self, text: str) -> str:
        """Remove null bytes and other postgres-breaking characters."""
        if not text:
            return ""
        return text.replace("\x00", "")
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts)
        except Exception as e:
            print(f"PDF Extraction error: {e}")
            return ""
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            return "\n\n".join(text_parts)
        except Exception as e:
            print(f"DOCX Extraction error: {e}")
            return ""
    
    async def _extract_metadata(
        self, 
        raw_text: str, 
        document_type: str,
        filename: str
    ) -> Dict[str, Any]:
        """Use LLM to extract structured metadata from document text."""
        
        # Truncate text if too long
        max_chars = 50000
        text_for_analysis = raw_text[:max_chars]
        
        # Build prompt
        prompt = f"""
        Extract structured metadata from this government contract document.
        Return ONLY valid JSON.
        
        DOCUMENT TEXT:
        {text_for_analysis}
        
        FIELDS TO EXTRACT:
        - contract_number (string)
        - contract_title (string)
        - customer_agency (string - e.g., "Department of Veterans Affairs", "US Coast Guard")
        - customer_command (string - specific command/facility)
        - contract_value (number or string with $ amount)
        - pop_start (YYYY-MM-DD or null)
        - pop_end (YYYY-MM-DD or null)
        - naics_code (string)
        - clearance_level (string)
        - fte_count (number or null)
        - geographic_scope (string)
        - scope_summary (string - 2-3 sentences describing the work)
        - key_capabilities (array of strings)
        - performance_ratings (CRITICAL - see below)
        - locations (array of strings)
        
        ## CRITICAL: CPARS / Performance Evaluation Extraction
        
        **CPARS documents contain a ratings table. You MUST extract ratings if present.**
        
        The table often loses formatting in PDF extraction and may look like "messy" text lines.
        LOOK FOR THESE PATTERNS:
        1. "Quality: Very Good" or "Quality.....Very Good"
        2. "Quality N/A Very Good" (The "N/A" is typical for "Past Rating")
        3. "QUALITY OF PRODUCT... Very Good"
        4. "Schedule . . . . . Satisfactory"
        
        **MANDATORY FIELDS TO HUNT FOR:**
        - Quality (of Product/Service)
        - Schedule (Timeliness)
        - Cost Control
        - Management (Business Relations)
        - Small Business (Subcontracting)
        - Regulatory Compliance
        
        Ratings are usually: "Exceptional", "Very Good", "Satisfactory", "Marginal", "Unsatisfactory".
        
        **REQUIRED: Extract into performance_ratings field as:**
        ```json
        {{
          "Quality": "Very Good",
          "Schedule": "Satisfactory", 
          "Management": "Exceptional",
          "Small_Business": "Marginal",
          "Recommendation": "Would recommend" 
        }}
        ```
        If a rating is "Marginal" or "Unsatisfactory", you MUST include it.
        
        ## CRITICAL: DOCUMENT CLASSIFICATION
        Analyze the document content to determine its true nature.
        - **CONTRACT_CPARS**: A formal contract, task order, or CPARS evaluation.
        - **NARRATIVE**: A capability statement, marketing slick, or general company overview (no specific contract #).
        - **SOLICITATION**: An RFP, PWS, or SOW (the requirements document).
        
        Return field: "actual_document_type": "CONTRACT_CPARS" | "NARRATIVE" | "SOLICITATION"
        
        ## JSON OUTPUT FORMAT:
        {{
            "contract_number": "...",
            "contract_title": "...",
            "actual_document_type": "CONTRACT_CPARS", 
            ...
            "performance_ratings": {{
                "Quality": "Very Good",
                "Schedule": "Satisfactory", 
                "Management": "Exceptional",
                "Small Business Subcontracting": "Marginal",
                "Regulatory Compliance": "Satisfactory"
            }}
        }}
        ```
        
        **Rating values are one of:** Exceptional, Very Good, Satisfactory, Marginal, Unsatisfactory, N/A
        
        **If ANY rating is "Marginal" or "Unsatisfactory", also include:**
        ```json
        {{
          "negative_ratings": ["Small Business Subcontracting: Marginal"]
        }}
        ```
        
        **If contractor disputes are mentioned, include:**
        ```json
        {{
          "contractor_disputes": "Contractor disputes Schedule rating..."
        }}
        ```
        
        If this is NOT a CPARS document, set performance_ratings to empty dict {{}}.
        
        Return ONLY the JSON object, no other text.
        """
        
        # Call LLM
        response = await self.llm.extract_metadata(prompt)
        
        # Ensure we always return a dict
        if not isinstance(response, dict):
            response = {}
        
        # Validate and enhance document type classification using heuristics
        llm_doc_type = response.get("actual_document_type", "").upper()
        heuristic_type = self._classify_document_type(filename, text_for_analysis, response)
        
        # Use heuristic if LLM classification is missing or uncertain
        if not llm_doc_type or llm_doc_type == "SOLICITATION" and document_type.lower() not in ["rfp", "solicitation", "pws", "sow"]:
            response["actual_document_type"] = heuristic_type
        else:
            response["actual_document_type"] = llm_doc_type
        
        # For CPARS/contract evaluation docs, add a second-pass heuristic parser
        # so that obvious ratings in the raw text are not missed when the LLM
        # fails to populate performance_ratings reliably.
        is_cpars_doc = (
            document_type.lower() == "cpars"
            or str(response.get("actual_document_type", "")).upper() == "CONTRACT_CPARS"
        )
        if is_cpars_doc:
            existing_ratings = response.get("performance_ratings") or {}
            if not existing_ratings:
                heuristic_ratings = self._extract_cpars_ratings_from_text(text_for_analysis)
                if heuristic_ratings:
                    response["performance_ratings"] = heuristic_ratings
        
        return response
    
    def _detect_mime_type(self, filename: str) -> str:
        """Detect mime type from filename."""
        ext = filename.lower().split(".")[-1] if "." in filename else ""
        mime_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc": "application/msword",
            "txt": "text/plain"
        }
        return mime_types.get(ext, "application/octet-stream")

    def _parse_currency(self, value: Any) -> float:
        if not value:
            return None
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, str):
            # Remove $ and ,
            cleaned = value.replace('$', '').replace(',', '').strip()
            try:
                return float(cleaned)
            except:
                return None
        return None

    def _parse_date(self, value: Any):
        if not value:
            return None
        try:
             # Parse and return date object
             parsed = datetime.strptime(str(value), "%Y-%m-%d")
             return parsed.date()  # Return date object, not string
        except:
            return None

    def _parse_int(self, value: Any) -> int:
        if not value:
            return None
        try:
            return int(value)
        except:
            return None
    
    def _parse_locations(self, value: Any):
        """Parse locations field - can be string or list."""
        if not value:
            return None
        if isinstance(value, list):
            return value
        if isinstance(value, str):
            # If it's a single string, return as single-item list
            return [value]
        return None

    def _get_cpars_rating(self, ratings: Dict[str, Any], possible_keys: List[str]) -> Optional[str]:
        """
        Safely extract a CPARS rating value from a flexible performance_ratings
        dict where keys may vary slightly (e.g. 'Quality of Product/Service').
        """
        if not ratings:
            return None
        
        # Normalize keys for case/spacing/underscore differences
        normalized: Dict[str, str] = {}
        for key, value in ratings.items():
            if key is None:
                continue
            norm_key = str(key).lower().replace("_", " ").strip()
            normalized[norm_key] = str(value)
        
        for key in possible_keys:
            lookup = key.lower().strip()
            if lookup in normalized and normalized[lookup]:
                return normalized[lookup]
        
        return None

    def _classify_document_type(self, filename: str, text: str, parsed_content: Dict[str, Any]) -> str:
        """
        Classify document type using filename patterns and content heuristics.
        Returns: "CONTRACT_CPARS", "NARRATIVE", or "SOLICITATION"
        """
        filename_lower = filename.lower()
        text_lower = text.lower()[:5000]  # Check first 5k chars for performance
        
        # Check filename patterns first (most reliable)
        if any(keyword in filename_lower for keyword in ["cpars", "past performance", "contract", "task order", "to-", "mod-"]):
            # But exclude if it's clearly a narrative
            if not any(keyword in filename_lower for keyword in ["capability", "narrative", "overview", "statement"]):
                return "CONTRACT_CPARS"
        
        if any(keyword in filename_lower for keyword in ["capability", "narrative", "overview", "statement", "company profile"]):
            return "NARRATIVE"
        
        if any(keyword in filename_lower for keyword in ["rfp", "solicitation", "pws", "sow", "amendment", "amendment"]):
            return "SOLICITATION"
        
        # Check content heuristics
        # Contract indicators
        contract_indicators = [
            "contract number",
            "task order",
            "modification",
            "period of performance",
            "cpars",
            "performance evaluation",
            "contractor performance",
        ]
        
        # Narrative indicators
        narrative_indicators = [
            "capability statement",
            "company overview",
            "our capabilities",
            "we provide",
            "about our company",
        ]
        
        # Solicitation indicators
        solicitation_indicators = [
            "section l",
            "section m",
            "instructions to offerors",
            "evaluation factors",
            "request for proposal",
        ]
        
        contract_score = sum(1 for indicator in contract_indicators if indicator in text_lower)
        narrative_score = sum(1 for indicator in narrative_indicators if indicator in text_lower)
        solicitation_score = sum(1 for indicator in solicitation_indicators if indicator in text_lower)
        
        # Check for contract number in parsed content
        has_contract_number = bool(parsed_content.get("contract_number"))
        
        # Decision logic
        if solicitation_score >= 2:
            return "SOLICITATION"
        elif has_contract_number and contract_score >= 2:
            return "CONTRACT_CPARS"
        elif narrative_score >= 2:
            return "NARRATIVE"
        elif has_contract_number:
            return "CONTRACT_CPARS"
        elif contract_score > narrative_score:
            return "CONTRACT_CPARS"
        elif narrative_score > 0:
            return "NARRATIVE"
        else:
            # Default based on document_type parameter
            return "CONTRACT_CPARS"  # Safe default - assume it's a contract unless proven otherwise
    
    def _extract_cpars_ratings_from_text(self, text: str) -> Dict[str, str]:
        """
        Heuristic CPARS rating extractor that works directly on the raw PDF text.
        
        Many CPARS PDFs lose their table formatting in extraction. This parser
        scans line-by-line for known rating categories and standard adjectival
        ratings (Exceptional, Very Good, Satisfactory, Marginal, Unsatisfactory, N/A).
        """
        if not text:
            return {}
        
        rating_words = [
            "Exceptional",
            "Very Good",
            "Satisfactory",
            "Marginal",
            "Unsatisfactory",
            "N/A",
            "NA",
        ]
        
        # Map of substrings to normalized keys
        label_map = {
            "quality of product/service": "Quality",
            "quality of product": "Quality",
            "quality": "Quality",
            "schedule": "Schedule",
            "cost control": "Cost",
            "cost": "Cost",
            "management/business relations": "Management",
            "business relations": "Management",
            "management": "Management",
            "small business subcontracting": "Small Business",
            "small business": "Small Business",
            "regulatory compliance": "Regulatory Compliance",
        }
        
        results: Dict[str, str] = {}
        lines = text.splitlines()
        
        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                continue
            
            lower_line = line.lower()
            
            # Match labels in order of specificity (longer/more specific first)
            sorted_labels = sorted(label_map.items(), key=lambda x: len(x[0]), reverse=True)
            
            for label_substr, norm_key in sorted_labels:
                if label_substr not in lower_line:
                    continue
                
                # Find the position of the label
                label_pos = lower_line.find(label_substr)
                
                # Find ratings near this label (within reasonable distance)
                # Look for ratings after the label
                text_after_label = lower_line[label_pos + len(label_substr):]
                
                # Try to find a rating word after the label
                found_rating: Optional[str] = None
                rating_positions = []
                for word in rating_words:
                    token = word.lower().replace("/", "")
                    pos = text_after_label.replace("/", "").find(token)
                    if pos != -1:
                        rating_positions.append((pos, word))
                
                if rating_positions:
                    # Use the closest (leftmost) rating after the label
                    rating_positions.sort(key=lambda x: x[0])
                    found_rating = "N/A" if rating_positions[0][1] in ("N/A", "NA") else rating_positions[0][1]
                    
                    # Only set if not already captured (first match wins)
                    if norm_key not in results:
                        results[norm_key] = found_rating
                    break  # Found a label+rating pair for this line, move to next line
        
        return results
